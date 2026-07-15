# -*- coding: utf-8 -*-
"""Build the PoC Word template `alimony-poc.docx`.

PoC scaffolding ONLY. In the target flow a lawyer authors this .docx directly in Word
(and a Word add-in writes the metadata blob). We generate it from code here so the
template is reviewable in git as a diff instead of an opaque binary — the output is a
real .docx that Word opens and edits normally.

Run:  python apps/doc-engine/templates/build_alimony_poc.py

Template conventions proven by the PoC (see fill.py):
- {%p ... %} paragraph tags must be alone in their paragraph — docxtpl removes the
  whole paragraph containing the tag.
- Values that engine B computes (declension, dates-in-words) are referenced by the
  naming contract in extract.py ENGINE_B_PROVIDED; everything else must be a form field.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).parent / "alimony-poc.docx"


def _p(doc, text, *, align=None, bold=False, size=12, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def build() -> Path:
    doc = Document()

    # ── шапка ────────────────────────────────────────────────────────────────
    _p(doc, "Позивач: {{ last_name }} {{ first_name }} {{ middle_name }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _p(doc, "дата народження: {{ birth_date_words }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _p(doc, "адреса реєстрації: {{ registered_address }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    _p(doc, "Відповідач: {{ defendant_last_name }} {{ defendant_first_name }} {{ defendant_middle_name }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _p(doc, "адреса реєстрації: {{ defendant_registered_address }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

    # ── заголовок ────────────────────────────────────────────────────────────
    _p(doc, "ПОЗОВНА ЗАЯВА", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=0)
    _p(doc, "про стягнення аліментів на утримання дитини", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    # ── обставини ────────────────────────────────────────────────────────────
    _p(doc, "На утриманні позивача перебувають діти:", space_after=0)
    _p(doc, "{{ children_details }}", space_after=12)

    _p(
        doc,
        "Починаючи з {{ abandonment_date_words }} відповідач не бере участі в утриманні "
        "дитини та не надає жодної матеріальної допомоги.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )

    # ── ветвление 1: працевлаштування відповідача (guard = defendant_employed, choice)
    _p(doc, "{%p if defendant_employed == 'yes' %}", space_after=0)
    _p(
        doc,
        "Відповідач офіційно працевлаштований: {{ defendant_employer }}. "
        "Отже, стягнення можливе за місцем отримання доходу.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )
    _p(doc, "{%p endif %}", space_after=0)

    # ── ветвление 2: форма стягнення (guard = alimony_type, choice) ──────────
    _p(doc, "{%p if alimony_type == 'fixed' %}", space_after=0)
    _p(
        doc,
        "Оскільки дохід відповідача є нерегулярним, прошу стягувати аліменти у твердій "
        "грошовій сумі — {{ alimony_fixed_amount }} грн щомісяця.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )
    _p(doc, "{%p else %}", space_after=0)
    _p(
        doc,
        "Прошу стягувати аліменти у частці від усіх видів заробітку (доходу) відповідача.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )
    _p(doc, "{%p endif %}", space_after=0)

    # ── прохальна частина ────────────────────────────────────────────────────
    _p(doc, "ПРОШУ:", bold=True, space_after=6)
    _p(
        doc,
        "Стягнути з {{ defendant_genitive }} на користь {{ plaintiff_genitive }} аліменти "
        "на утримання дитини, починаючи з {{ alimony_start_date_words }} і до досягнення "
        "дитиною повноліття.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=24,
    )

    # ── подпись ──────────────────────────────────────────────────────────────
    _p(doc, "{{ alimony_start_date_words }}", space_after=0)
    _p(doc, "{{ last_name }} {{ first_name }} {{ middle_name }}     _______________", space_after=0)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"wrote {path}")
