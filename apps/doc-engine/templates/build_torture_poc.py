# -*- coding: utf-8 -*-
"""Build `torture-poc.docx` — the adversarial Word template.

WHY THIS EXISTS: the Word add-in spike needs a baseline. This template deliberately packs
every authoring gotcha we have already paid for, so "the add-in didn't break anything" is a
measurable claim instead of a feeling. The document is FICTIONAL — no legal review needed,
which is the point: we can make it as nasty as we like.

Gotchas encoded here (each one cost us a session to find):
  1. {%p ... %} paragraph tags must be ALONE in their paragraph (docxtpl deletes the whole
     paragraph holding the tag).                                              — s84
  2. TWO ADJACENT branches back-to-back: a naive "wrap the selection" add-in that reuses one
     insertion point will interleave them.                                    — new
  3. A BARE `{% if has_contract %}` boolean guard, no comparison. jinja2's find_all() walks
     descendants only, so the bare Name node is invisible unless collected explicitly. This
     exact shape was a real bug found by a negative test.                      — s85
  4. NESTED if inside an if/else — the pair-matching an add-in must never get wrong.
  5. {%tr%} table loop: for/endfor in SEPARATE control ROWS, content between. Both in one row
     destroys the row.                                                        — s84
  6. An OPTIONAL field guarded by its own presence — the shape that would have prevented
     "________ року народження" leaking into ПРОШУ СУД.        — demo D7, finding A
  7. Declension of BOTH genders incl. a feminine indeclinable surname (Петренко stays,
     Ковальчук declines). pymorphy parse[0] without gender fails silently.     — s84
  8. Dates in words, incl. a date that only exists on one branch.
  9. A variable in the HEADER — declared_variables reads headers/footers, not just the body.
 10. Signature block with keep_with_next so it cannot be orphaned across a page break. — s84

Run:  python apps/doc-engine/templates/build_torture_poc.py
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))
from doc_engine.docx_meta import embed_in_docx  # noqa: E402

OUT = Path(__file__).parent / "torture-poc.docx"
META = Path(__file__).parent / "torture-poc.meta.json"


def _p(doc, text, *, align=None, bold=False, size=12, space_after=6, keep=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if keep:
        p.paragraph_format.keep_with_next = True
    return p


def _cell(cell, text, *, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"


def build() -> Path:
    doc = Document()

    # ── gotcha 9: a variable in the HEADER ───────────────────────────────────
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.text = "Справа № {{ case_ref }}"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ── шапка ────────────────────────────────────────────────────────────────
    _p(doc, "Позивач: {{ last_name }} {{ first_name }} {{ middle_name }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _p(doc, "дата народження: {{ birth_date_words }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _p(doc, "адреса: {{ registered_address }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    _p(doc, "Відповідач: {{ defendant_last_name }} {{ defendant_first_name }} {{ defendant_middle_name }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _p(doc, "адреса: {{ defendant_registered_address }}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    _p(doc, "ТЕСТОВА ЗАЯВА", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=0)
    _p(doc, "про стягнення боргу (вигаданий документ)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    # ── gotcha 3 + 4: bare boolean guard, with a NESTED if inside an if/else ──
    _p(doc, "{%p if claim_type == 'main' %}", space_after=0)
    _p(
        doc,
        "Це основна вимога. Відповідач отримав від позивача грошові кошти у розмірі "
        "{{ debt_amount }} грн та не повернув їх у встановлений строк.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )
    _p(doc, "{%p if has_contract %}", space_after=0)
    _p(
        doc,
        "Між сторонами укладено письмовий договір позики від {{ contract_date_words }}, "
        "що підтверджує факт передачі коштів.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )
    _p(doc, "{%p endif %}", space_after=0)
    _p(doc, "{%p else %}", space_after=0)
    _p(
        doc,
        "Це додаткова вимога, похідна від основного зобов'язання на суму {{ debt_amount }} грн.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )
    _p(doc, "{%p endif %}", space_after=0)

    # ── gotcha 2: TWO ADJACENT branches, back-to-back, no text between ────────
    _p(doc, "{%p if interest_claimed %}", space_after=0)
    _p(
        doc,
        "Окрім суми основного боргу, позивач просить стягнути проценти за користування "
        "чужими грошовими коштами.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )
    _p(doc, "{%p endif %}", space_after=0)
    # gotcha 6: presence-guard on an OPTIONAL field — the shape that stops "________"
    _p(doc, "{%p if defendant_birth_date_words %}", space_after=0)
    # NB: date_words() already ends in "року", so the template must NOT add it again —
    # "{{ x_words }} року народження" double-prints. Engine A's formatDate gives "03.11.1988"
    # where that same phrasing reads correctly, so this trap is invisible to a lawyer who
    # only sees the field NAME in the pane. Hence: the add-in must preview the VALUE.
    _p(
        doc,
        "Відповідач, {{ defendant_birth_date_words }} народження, є повнолітнім та "
        "дієздатним.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )
    _p(doc, "{%p endif %}", space_after=0)

    # ── gotcha 5: {%tr%} loop — for/endfor in SEPARATE control rows ───────────
    _p(doc, "Додатки:", bold=True, space_after=6)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    _cell(table.rows[0].cells[0], "№", bold=True)
    _cell(table.rows[0].cells[1], "Найменування доказу", bold=True)
    _cell(table.rows[1].cells[0], "{%tr for item in evidence %}")
    _cell(table.rows[2].cells[0], "{{ loop.index }}")
    _cell(table.rows[2].cells[1], "{{ item }}")
    _cell(table.rows[3].cells[0], "{%tr endfor %}")

    _p(doc, "", space_after=12)

    # ── прохальна частина: gotcha 7 — both genders declined ──────────────────
    _p(doc, "ПРОШУ:", bold=True, space_after=6)
    _p(
        doc,
        "Стягнути з {{ defendant_genitive }} на користь {{ plaintiff_genitive }} суму боргу "
        "у розмірі {{ debt_amount }} грн.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=24,
    )

    # ── gotcha 10: signature must not be orphaned ─────────────────────────────
    _p(doc, "{{ filing_date_words }}", space_after=0, keep=True)
    _p(doc, "{{ last_name }} {{ first_name }} {{ middle_name }}     _______________", space_after=0)

    doc.save(OUT)
    # The .docx is the artifact: metadata lives INSIDE it, exactly as the add-in will leave it.
    embed_in_docx(OUT, json.loads(META.read_text(encoding="utf-8")))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"wrote {path} (metadata embedded in customXml)")
