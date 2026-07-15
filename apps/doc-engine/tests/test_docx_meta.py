# -*- coding: utf-8 -*-
"""Tests for service metadata embedded inside the .docx (customXml).

This is the store the future Word add-in writes through, so the invariants that matter are:
the blob survives a round-trip byte-for-byte in meaning, re-embedding never accumulates parts,
a foreign custom XML part (python-docx's own default template ships one!) is never mistaken
for ours or clobbered, and the surgery leaves a package Word/LibreOffice can still open.
"""
from __future__ import annotations

import json
import zipfile

import pytest
from docx import Document
from docxtpl import DocxTemplate

from doc_engine.docx_meta import NS, embed_in_docx, read_from_docx
from doc_engine.extract import extract, load_metadata

META = {
    "service": {"slug": "svc", "title": "Тест"},
    "tabs": [{"id": "main", "label": "Головна"}],
    "fields": [{"id": "last_name", "tab": "main", "type": "text", "label": "Прізвище"}],
}


def _docx(tmp_path, text="{{ last_name }}"):
    doc = Document()
    doc.add_paragraph(text)
    path = tmp_path / "t.docx"
    doc.save(path)
    return path


def _custom_parts(path):
    with zipfile.ZipFile(path) as z:
        return [n for n in z.namelist() if n.startswith("customXml/item") and "Props" not in n]


def test_read_returns_none_when_absent(tmp_path):
    assert read_from_docx(_docx(tmp_path)) is None


def test_round_trip_preserves_blob(tmp_path):
    p = _docx(tmp_path)
    embed_in_docx(p, META)
    assert read_from_docx(p) == META


def test_round_trip_preserves_cyrillic_and_nested(tmp_path):
    p = _docx(tmp_path)
    meta = json.loads(json.dumps(META))
    meta["fields"][0]["hint"] = "Лише кирилиця, апостроф'"
    meta["fields"][0]["show_if"] = {"field": "x", "operator": "==", "value": "так"}
    embed_in_docx(p, meta)
    assert read_from_docx(p) == meta


def test_reembedding_is_idempotent(tmp_path):
    """Re-saving a service must not accumulate a new part each time."""
    p = _docx(tmp_path)
    embed_in_docx(p, META)
    after_first = _custom_parts(p)
    embed_in_docx(p, {**META, "service": {"slug": "svc", "title": "Змінено"}})
    embed_in_docx(p, META)
    assert _custom_parts(p) == after_first
    assert read_from_docx(p) == META


def test_foreign_custom_xml_part_is_not_ours(tmp_path):
    """python-docx's default template already ships customXml/item1.xml — never claim it."""
    p = _docx(tmp_path)
    foreign = [n for n in _custom_parts(p)]
    embed_in_docx(p, META)
    with zipfile.ZipFile(p) as z:
        for n in foreign:
            assert NS not in z.read(n).decode("utf-8", "ignore"), f"{n} must stay untouched"
    assert read_from_docx(p) == META


def test_package_still_opens_and_renders(tmp_path):
    """Zip surgery must not corrupt the OPC package."""
    p = _docx(tmp_path)
    embed_in_docx(p, META)
    assert DocxTemplate(str(p)).get_undeclared_template_variables() == {"last_name"}
    Document(str(p))  # raises if the package is malformed


def test_embed_to_separate_destination(tmp_path):
    src = _docx(tmp_path)
    dst = tmp_path / "out.docx"
    embed_in_docx(src, META, dst)
    assert read_from_docx(dst) == META
    assert read_from_docx(src) is None  # source untouched


def test_load_metadata_prefers_embedded_over_sidecar(tmp_path):
    p = _docx(tmp_path)
    embed_in_docx(p, META)
    sidecar = {**META, "service": {"slug": "stale", "title": "Застарілий sidecar"}}
    (tmp_path / "t.meta.json").write_text(json.dumps(sidecar, ensure_ascii=False), encoding="utf-8")

    meta, origin = load_metadata(p)
    assert origin == "embedded"
    assert meta["service"]["slug"] == "svc"
    assert extract(p)["metadata_origin"] == "embedded"


def test_load_metadata_falls_back_to_sidecar(tmp_path):
    p = _docx(tmp_path)
    (tmp_path / "t.meta.json").write_text(json.dumps(META, ensure_ascii=False), encoding="utf-8")
    meta, origin = load_metadata(p)
    assert origin == "sidecar"
    assert meta == META


def test_load_metadata_errors_when_nothing_anywhere(tmp_path):
    with pytest.raises(ValueError, match="no metadata"):
        load_metadata(_docx(tmp_path))
