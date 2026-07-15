# -*- coding: utf-8 -*-
"""Tests for the Word→service extractor.

Regression origin: guard_variables missed a bare `{% if flag %}` because jinja2's
find_all() walks descendants only — the test node IS the Name there, so it was never
collected. Compound tests (`{% if x == 'y' %}`) hid the bug: their Name sits inside a
Compare. A missed guard means the bridge cannot enforce "branch only on askable data",
so a document could branch on something the interview never asks.
"""
from __future__ import annotations

import json

import pytest
from docx import Document
from docxtpl import DocxTemplate

from doc_engine.extract import derived_sources, extract, guard_variables


def _docx(tmp_path, *paragraphs):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    path = tmp_path / "t.docx"
    doc.save(path)
    return DocxTemplate(str(path))


def test_guard_bare_name(tmp_path):
    """`{% if flag %}` — the test node itself is a Name (the regression)."""
    tpl = _docx(tmp_path, "{%p if has_children %}", "текст", "{%p endif %}")
    assert guard_variables(tpl) == {"has_children"}


def test_guard_comparison(tmp_path):
    """`{% if x == 'y' %}` — Name nested inside a Compare."""
    tpl = _docx(tmp_path, "{%p if alimony_type == 'fixed' %}", "текст", "{%p endif %}")
    assert guard_variables(tpl) == {"alimony_type"}


def test_guard_elif_and_and(tmp_path):
    tpl = _docx(
        tmp_path,
        "{%p if a == 'x' and b %}", "1",
        "{%p elif c %}", "2",
        "{%p endif %}",
    )
    assert guard_variables(tpl) == {"a", "b", "c"}


def test_guard_none_when_no_branching(tmp_path):
    tpl = _docx(tmp_path, "Просто {{ last_name }} без розгалужень")
    assert guard_variables(tpl) == set()


def test_derived_dates_and_declension():
    d = derived_sources({
        "birth_date": "date",
        "last_name": "text", "first_name": "text", "middle_name": "text",
        "note": "textarea",
    })
    assert d["birth_date_words"] == ["birth_date"]
    assert d["plaintiff_genitive"] == ["last_name", "first_name", "middle_name"]
    assert "note_words" not in d          # only date fields get _words
    assert "defendant_genitive" not in d  # form carries no defendant name fields


def test_extract_reads_template_and_metadata(tmp_path):
    """extract() returns variables incl. guards, plus the lawyer-authored metadata."""
    doc = Document()
    doc.add_paragraph("{{ last_name }} — {{ birth_date_words }}")
    doc.add_paragraph("{%p if flag %}")
    doc.add_paragraph("{{ note }}")
    doc.add_paragraph("{%p endif %}")
    tpl_path = tmp_path / "svc.docx"
    doc.save(tpl_path)

    meta = {
        "service": {"slug": "svc", "title": "T"},
        "tabs": [{"id": "main", "label": "M"}],
        "fields": [
            {"id": "last_name", "tab": "main", "type": "text", "label": "Прізвище"},
            {"id": "birth_date", "tab": "main", "type": "date", "label": "Дата"},
            {"id": "flag", "tab": "main", "type": "boolean", "label": "Прапорець"},
            {"id": "note", "tab": "main", "type": "textarea", "label": "Нотатка"},
        ],
    }
    (tmp_path / "svc.meta.json").write_text(json.dumps(meta, ensure_ascii=False), encoding="utf-8")

    got = extract(tpl_path)
    assert got["guards"] == ["flag"]
    assert "birth_date_words" in got["variables"]
    assert got["derived"]["birth_date_words"] == ["birth_date"]
    assert got["metadata"]["service"]["slug"] == "svc"


def test_extract_rejects_metadata_missing_keys(tmp_path):
    doc = Document()
    doc.add_paragraph("{{ x }}")
    tpl_path = tmp_path / "svc.docx"
    doc.save(tpl_path)
    (tmp_path / "svc.meta.json").write_text(json.dumps({"service": {}}), encoding="utf-8")

    with pytest.raises(ValueError, match="missing required key"):
        extract(tpl_path)
