# -*- coding: utf-8 -*-
"""Baseline for the Word add-in spike: the adversarial template, rendered.

This is the ONLY thing that makes "the add-in didn't break anything" measurable. Every
assertion here pins a gotcha that cost us a session to find (see build_torture_poc.py).
When the add-in has touched torture-poc.docx and saved it from a live Word, this suite
must still pass byte-for-behaviour — otherwise Word's round-trip destroyed something.
"""
from __future__ import annotations

import io
import json
from pathlib import Path

import pytest
from docx import Document

from doc_engine import MissingFieldsError, build_context, extract, render_to_bytes

TEMPLATES = Path(__file__).resolve().parents[1] / "templates"
TPL = TEMPLATES / "torture-poc.docx"

MAXIMAL = dict(
    case_ref="123/4567/26",
    filing_date="2026-07-15",
    last_name="Петренко", first_name="Олена", middle_name="Іванівна",
    birth_date="1990-05-12", registered_address="м. Київ, вул. Хрещатик, 1",
    defendant_last_name="Ковальчук", defendant_first_name="Ігор",
    defendant_middle_name="Васильович",
    defendant_registered_address="м. Львів, вул. Січових Стрільців, 5",
    defendant_birth_date="1988-11-03",
    has_contract=True, contract_date="2024-03-10", debt_amount="150000",
    claim_type="main", interest_claimed=True,
    evidence=["Копія договору позики", "Розписка про отримання коштів", "Банківська виписка"],
)

#: Every branch OFF, every optional field blank — the shape that made engine A print "________".
MINIMAL = dict(
    MAXIMAL,
    has_contract=False, contract_date=None, defendant_birth_date=None,
    claim_type="additional", interest_claimed=False,
    evidence=["Листування сторін"],
)


@pytest.fixture(scope="module")
def meta() -> dict:
    return json.loads((TEMPLATES / "torture-poc.meta.json").read_text(encoding="utf-8"))


def _render(meta: dict, answers: dict):
    ctx, flags = build_context(meta, answers)
    doc = Document(io.BytesIO(render_to_bytes(str(TPL), ctx)))
    body = "\n".join(p.text for p in doc.paragraphs)
    header = doc.sections[0].header.paragraphs[0].text
    rows = [[c.text for c in r.cells] for t in doc.tables for r in t.rows]
    return body, header, rows, flags


# ── gotcha 9: variables in the header are real variables ─────────────────────
def test_header_variable_is_declared_and_rendered(meta):
    assert "case_ref" in extract(TPL)["variables"]
    _, header, _, _ = _render(meta, MAXIMAL)
    assert header == "Справа № 123/4567/26"


# ── gotcha 3: a BARE boolean guard is seen by the AST walk ───────────────────
def test_bare_boolean_guards_are_detected(meta):
    guards = extract(TPL)["guards"]
    assert "has_contract" in guards          # bare `{% if has_contract %}`
    assert "interest_claimed" in guards      # bare, adjacent branch
    assert "claim_type" in guards            # compound `{% if claim_type == 'main' %}`


# ── gotchas 1, 2, 4: adjacent branches + nested if/else resolve independently ─
def test_maximal_branches_all_present(meta):
    body, _, _, _ = _render(meta, MAXIMAL)
    assert "Це основна вимога" in body
    assert "письмовий договір позики від 10 березня 2024 року" in body   # nested if
    assert "проценти за користування" in body                            # adjacent branch 1
    assert "3 листопада 1988 року народження" in body                    # adjacent branch 2
    assert "Це додаткова вимога" not in body                             # else arm gone


def test_minimal_branches_all_absent(meta):
    body, _, _, _ = _render(meta, MINIMAL)
    assert "Це додаткова вимога" in body
    assert "Це основна вимога" not in body
    assert "договір позики від" not in body      # nested if inside the dead arm
    assert "проценти за користування" not in body
    assert "народження, є повнолітнім" not in body


# ── gotcha 6 / demo-D7 finding A: a blank optional must not leak a placeholder ─
def test_blank_optional_leaves_no_hole(meta):
    body, _, _, _ = _render(meta, MINIMAL)
    assert "________ року народження" not in body
    assert "________ народження" not in body
    # The only underscores allowed are the signature rule the lawyer drew on purpose.
    lines_with_underscores = [ln for ln in body.splitlines() if "____" in ln]
    assert lines_with_underscores == ["Петренко Олена Іванівна     _______________"]


# ── gotcha 5: {%tr%} loop multiplies rows, control rows vanish ───────────────
def test_table_loop_expands_and_control_rows_disappear(meta):
    _, _, rows, _ = _render(meta, MAXIMAL)
    assert rows[0] == ["№", "Найменування доказу"]
    assert [r[1] for r in rows[1:]] == [
        "Копія договору позики", "Розписка про отримання коштів", "Банківська виписка",
    ]
    assert [r[0] for r in rows[1:]] == ["1", "2", "3"]   # loop.index
    assert not any("{%tr" in c for r in rows for c in r)


def test_table_loop_single_item(meta):
    _, _, rows, _ = _render(meta, MINIMAL)
    assert len(rows) == 2                                # header + one evidence row
    assert rows[1] == ["1", "Листування сторін"]


# ── gotcha 7: declension, both genders, no silent fallback ──────────────────
def test_declension_both_genders(meta):
    body, _, _, flags = _render(meta, MAXIMAL)
    # feminine + indeclinable surname: Петренко must NOT change, given name must
    assert "на користь Петренко Олени Іванівни" in body
    # masculine: every part declines
    assert "Стягнути з Ковальчука Ігоря Васильовича" in body
    assert flags == []


# ── gotcha 8: dates in words, incl. one that only exists on a branch ─────────
def test_dates_in_words(meta):
    body, _, _, _ = _render(meta, MAXIMAL)
    assert "дата народження: 12 травня 1990 року" in body
    assert "15 липня 2026 року" in body                  # filing date, above the signature


# ── the anti-________ contract: a REQUIRED field with no answer must refuse ──
def test_missing_required_field_refuses_to_render(meta):
    answers = dict(MAXIMAL)
    del answers["debt_amount"]
    ctx, _ = build_context(meta, answers)
    with pytest.raises(MissingFieldsError) as e:
        render_to_bytes(str(TPL), ctx)
    assert "debt_amount" in e.value.fields


# ── gotcha 10: the signature block must not be orphaned ─────────────────────
def test_signature_paragraph_keeps_with_next():
    doc = Document(str(TPL))
    sig = [p for p in doc.paragraphs if "filing_date_words" in p.text]
    assert sig and sig[0].paragraph_format.keep_with_next is True
